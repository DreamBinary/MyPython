
from docx import Document
doc=Document(r"D:\Users\16477\PycharmProjects\pythonProject\word\猪猪.docx")
print(doc.paragraphs)
for paragraph in doc.paragraphs:
    print(paragraph.text)
print(dir(Document()))


























