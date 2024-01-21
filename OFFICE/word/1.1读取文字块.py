from docx import Document
doc=Document(r"D:\Users\16477\PycharmProjects\pythonProject\word\猪猪.docx")
print(doc.paragraphs)
print('---------------------')
paragraph=doc.paragraphs[0]
runs=paragraph.runs
print(runs)
print('---------------------')
for run in paragraph.runs:
    print(run.text)
#一个run对象是相同样式文本的延续 文字块
print('---------------------')


























