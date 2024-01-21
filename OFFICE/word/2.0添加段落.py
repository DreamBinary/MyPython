from docx import Document
doc=Document(r"D:\Users\16477\PycharmProjects\pythonProject\word\猪猪.docx")
#doc.add_heading("heading",level=1)
par1=doc.add_paragraph("加一段落")
#赋值便于后面进行格式调整
doc.add_paragraph("再加一段")
doc.save(r"D:\Users\16477\PycharmProjects\pythonProject\word\添加段落.docx")

























