from docx import Document
doc=Document(r"D:\Users\16477\PycharmProjects\pythonProject\word\猪猪.docx")
par=doc.add_paragraph()
par.add_run("我是加粗文字块").bold=True
par.add_run("我是正常文字块")
par.add_run("我是斜体文字块").italic=True
par.add_run("我是红色文字块").color='FF0000' #不红
doc.save(r"D:\Users\16477\PycharmProjects\pythonProject\word\添加文字块.docx")


























