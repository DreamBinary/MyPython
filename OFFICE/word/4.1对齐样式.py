from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
doc=Document(r"D:\Users\16477\PycharmProjects\pythonProject\word\猪猪.docx")
print(doc.paragraphs[0].text)
doc.paragraphs[0].alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
#居中对齐 LEFT CENTER RIGHT JUSTIFY DOSTRIBUTE JUSTIFY_MED JUSTIFT_HI JUSTIFY_LOW THAI_JUSTIFY
doc.save(r"D:\Users\16477\PycharmProjects\pythonProject\word\对齐样式.docx")





























