import random
import docx
full_text = []
doc = docx.Document('D:\\我好帅.docx')
doc_new = docx.Document()
doc_new = doc
paras = doc.paragraphs
for p in paras:
    full_text.append(p.text)
random.shuffle(full_text)
for i in range(len(full_text)):
    doc_new.paragraphs[i].text = full_text[i]
doc_new.save('D:\\我好帅2.0.docx')
